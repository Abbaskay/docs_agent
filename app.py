import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import io
import re
import json
import hashlib
import uuid
import time
import threading
import logging
import traceback
from datetime import datetime, timezone

from flask import Flask, render_template, request, jsonify, Response, stream_with_context
from flask_socketio import SocketIO, emit
from flask_cors import CORS

import pypdf
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

from config import DOCUMENT_AGENT_CONFIG
from core.memory import Memory
from core.runner import run_agent

from core.security import (
    sanitize_text,
    sanitize_filename,
    sanitize_html_content,
    validate_prompt,
    validate_document_text,
    validate_upload,
    MAX_DOCUMENT_LENGTH,
    MAX_PROMPT_LENGTH,
    strip_html_tags,
)
from core.rate_limiter import (
    generation_rate_limit,
    export_rate_limit,
    upload_rate_limit,
    burst_limiter,
    rate_limit_key,
    get_client_ip,
)
from core.database import init_db, db, Generation, Export, Upload, ApiUsage
from core.cache import cache_get, cache_set, cache_get_json, cache_set_json, make_cache_key
from core.exceptions import (
    AppError,
    ValidationError,
    FileValidationError,
    AIError,
    ExportError,
    RateLimitError,
    NotFoundError,
)

logging.basicConfig(
    level=getattr(logging, os.getenv("LOG_LEVEL", "INFO")),
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("aidoc")

app = Flask(__name__)
app.secret_key = os.getenv(
    "FLASK_SECRET_KEY",
    os.urandom(24).hex(),
)

ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "*")
CORS(
    app,
    origins=ALLOWED_ORIGINS.split(","),
    supports_credentials=True,
    methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Content-Type", "X-Session-Id", "X-API-Key", "X-CSRF-Token"],
)

socketio = SocketIO(
    app,
    cors_allowed_origins=ALLOWED_ORIGINS.split(","),
    async_mode="threading",
    ping_timeout=int(os.getenv("SOCKET_PING_TIMEOUT", "60")),
    ping_interval=int(os.getenv("SOCKET_PING_INTERVAL", "25")),
)

db = init_db(app)

MAX_DOCUMENT_CHARS = int(os.getenv("MAX_DOCUMENT_CHARS", "50000"))
DEFAULT_FILENAME = "my_document.txt"
FONT_DIR = Path(__file__).resolve().parent / "fonts"

sessions: dict[str, dict] = {}
_sessions_lock = threading.Lock()
_export_locks: dict[str, threading.Lock] = {}


def _get_export_lock(key: str) -> threading.Lock:
    if key not in _export_locks:
        _export_locks[key] = threading.Lock()
    return _export_locks[key]


def get_session(sid: str) -> dict:
    with _sessions_lock:
        if sid not in sessions:
            sessions[sid] = {
                "memory": Memory(),
                "thinking": [],
                "last_doc_hash": None,
                "doc_loaded": False,
                "last_download_name": "document",
                "created_at": time.time(),
            }
        return sessions[sid]


def cleanup_stale_sessions(max_age: int = 3600):
    now = time.time()
    with _sessions_lock:
        stale = [sid for sid, s in sessions.items() if now - s.get("created_at", 0) > max_age]
        for sid in stale:
            del sessions[sid]
    if stale:
        logger.info("Cleaned up %d stale sessions", len(stale))


threading.Thread(target=cleanup_stale_sessions, args=(3600,), daemon=True).start()


def document_hash(text: str) -> str | None:
    if not text:
        return None
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def safe_filename(value: str, default: str = "document") -> str:
    stem = re.sub(r"\.[A-Za-z0-9]{1,8}$", "", str(value).strip() or default)
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._-")
    return (stem or default)[:80]


def extract_text_from_file(uploaded_file) -> tuple[str, str | None]:
    filename = uploaded_file.filename.lower() if hasattr(uploaded_file, "filename") else ""
    try:
        uploaded_file.seek(0)
    except Exception:
        pass

    if filename.endswith(".pdf"):
        try:
            reader = pypdf.PdfReader(uploaded_file)
            if getattr(reader, "is_encrypted", False):
                try:
                    reader.decrypt("")
                except Exception:
                    return "", "Encrypted PDFs are not supported."
            pages = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    pages.append(page_text.strip())
            if not pages:
                return "", "No selectable text was found in this PDF."
            return "\n\n".join(pages), None
        except Exception as exc:
            return "", f"Could not read PDF: {exc}"

    if filename.endswith(".docx"):
        try:
            document = docx.Document(uploaded_file)
            parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            if not parts:
                return "", "No readable text was found in this DOCX."
            return "\n".join(parts), None
        except Exception as exc:
            return "", f"Could not read DOCX: {exc}"

    try:
        raw = uploaded_file.read()
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return raw.decode(encoding), None
            except UnicodeDecodeError:
                continue
        return "", "Unsupported text encoding."
    except Exception as exc:
        return "", f"Could not read file: {exc}"


def create_word_doc(text: str) -> bytes:
    content = str(text or "No document content was generated.")
    document = docx.Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
        elif re.fullmatch(r"[═─\-]{4,}", line):
            continue
        elif line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.isupper() and len(line) <= 80 and any(ch.isalpha() for ch in line):
            document.add_heading(line.title(), level=2)
        elif line.startswith(("• ", "- ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_pdf(text: str) -> bytes:
    content = str(text or "No document content was generated.")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    font_path = FONT_DIR / "ArialUnicode.ttf"
    font_name = "ArialUnicode" if font_path.exists() else "Helvetica"
    if font_path.exists():
        pdf.add_font(font_name, "", str(font_path), uni=True)
    pdf.set_font(font_name, "", 10)
    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            pdf.ln(4)
        elif re.fullmatch(r"[═─\-]{4,}", stripped):
            pdf.ln(2)
        elif stripped.isupper() and len(stripped) <= 80 and any(ch.isalpha() for ch in stripped):
            pdf.set_font(font_name, "", 13)
            pdf.cell(page_w, 8, stripped.title(), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font(font_name, "", 10)
        elif stripped.startswith(("• ", "- ")):
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(page_w - 5, 5, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(page_w - 5, 5, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(page_w, 5, stripped)

    return pdf.output()


def build_structured_log(entry: dict):
    try:
        now = datetime.now(timezone.utc).isoformat()
        entry["@timestamp"] = now
        logger.info("EVENT: %s", json.dumps(entry))
    except Exception:
        pass


def safe_agent_call(
    task: str,
    config,
    memory: Memory,
    on_tool_call=None,
) -> dict:
    start = time.time()
    try:
        result = run_agent(
            task=task,
            config=config,
            memory=memory,
            on_tool_call=on_tool_call,
        )
        duration_ms = int((time.time() - start) * 1000)
        result["_duration_ms"] = duration_ms
        return result
    except Exception as e:
        duration_ms = int((time.time() - start) * 1000)
        logger.error("Agent call failed: %s", traceback.format_exc())
        return {
            "answer": "I encountered an unexpected error while processing your request. Please try again.",
            "success": False,
            "trace": [],
            "iterations": 0,
            "tool_call_count": 0,
            "_duration_ms": duration_ms,
            "_error": str(e),
        }


def build_generation_key(session_id: str) -> str:
    return f"gen:{session_id}"


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------


@app.route("/")
def index():
    return render_template("index.html")


@app.errorhandler(404)
def not_found(e):
    return jsonify({"error": "Not found"}), 404


@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({"error": "Method not allowed"}), 405


@app.errorhandler(500)
def internal_error(e):
    build_structured_log({"event": "internal_error", "error": str(e)})
    return jsonify({"error": "Internal server error"}), 500


@app.errorhandler(AppError)
def handle_app_error(e: AppError):
    return jsonify(e.to_dict()), e.status


@app.before_request
def reject_large_payloads():
    cl = request.content_length
    if cl and cl > 5 * 1024 * 1024:
        return jsonify({"error": "Request too large"}), 413


# ---------------------------------------------------------------------------
# CHAT
# ---------------------------------------------------------------------------


@app.route("/api/chat", methods=["POST"])
@generation_rate_limit
def chat():
    data = request.get_json(silent=True)
    if not data:
        raise ValidationError("Request body must be valid JSON.")

    sid = data.get("session_id", request.remote_addr)
    raw_input = data.get("message", "")
    raw_doc_text = data.get("document_text", "")
    doc_filename = data.get("document_filename", DEFAULT_FILENAME)

    if not raw_input or not raw_input.strip():
        raise ValidationError("Message is required.")

    try:
        user_input = sanitize_text(strip_html_tags(raw_input), max_length=5000)
    except Exception:
        raise ValidationError("Invalid message format.")

    if not user_input.strip():
        raise ValidationError("Message cannot be empty.")

    if user_input != raw_input.strip():
        pass

    doc_text = ""
    if raw_doc_text and raw_doc_text.strip():
        try:
            doc_text = validate_document_text(raw_doc_text)
        except ValueError as e:
            raise ValidationError(str(e))

    doc_filename = sanitize_filename(doc_filename) or DEFAULT_FILENAME

    sess = get_session(sid)
    current_hash = document_hash(doc_text) if doc_text else None
    if current_hash and sess["last_doc_hash"] != current_hash:
        sess["doc_loaded"] = False

    should_mark_doc = bool(doc_text and not sess["doc_loaded"])
    if should_mark_doc:
        if len(doc_text) > MAX_DOCUMENT_CHARS:
            doc_text = doc_text[:MAX_DOCUMENT_CHARS]
        task = f"[DOCUMENT: {doc_filename}]\n{doc_text}\n\nRequest: {user_input}"
    else:
        task = user_input

    sess["thinking"] = []

    def on_tool_call(tool_name, tool_input, tool_output):
        sess["thinking"].append({
            "tool": tool_name,
            "input": tool_input,
            "output": str(tool_output)[:300],
        })

    gen_key = build_generation_key(sid)
    if not burst_limiter.try_acquire(gen_key):
        return jsonify({
            "error": "A generation is already in progress. Please wait.",
            "answer": "Please wait for the current generation to complete.",
            "thinking": [],
            "success": False,
            "session_id": sid,
        }), 429

    try:
        result = safe_agent_call(
            task=task,
            config=DOCUMENT_AGENT_CONFIG,
            memory=sess["memory"],
            on_tool_call=on_tool_call,
        )

        answer = result.get("answer", "Sorry, I could not complete that request.")
        success = result.get("success", False)

        gen_record = Generation(
            session_id=sid,
            doc_type="chat",
            prompt_length=len(task),
            status="completed" if success else "failed",
            success=success,
            error=result.get("_error") if not success else None,
            model=getattr(DOCUMENT_AGENT_CONFIG, "model", "deepseek-chat"),
            duration_ms=result.get("_duration_ms"),
            tool_calls=result.get("tool_call_count", 0),
        )
        try:
            db.session.add(gen_record)
            db.session.commit()
        except Exception:
            db.session.rollback()

        if should_mark_doc and success:
            sess["last_doc_hash"] = current_hash
            sess["doc_loaded"] = True
        sess["last_download_name"] = doc_filename if doc_text else "document_latest"

        resp = jsonify({
            "answer": answer,
            "thinking": sess["thinking"],
            "success": success,
            "session_id": sid,
        })
        return resp

    finally:
        burst_limiter.release(gen_key)


# ---------------------------------------------------------------------------
# UPLOAD
# ---------------------------------------------------------------------------


@app.route("/api/upload", methods=["POST"])
@upload_rate_limit
def upload():
    if "file" not in request.files:
        raise ValidationError("No file provided.")
    file = request.files["file"]
    if not file.filename or file.filename.strip() == "":
        raise ValidationError("Empty filename.")

    raw_filename = sanitize_filename(file.filename)
    raw_data = file.read()

    if not raw_data:
        raise FileValidationError("Empty file.")

    content_type = request.headers.get("Content-Type", "")
    mime_type = file.content_type or ""

    sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"

    try:
        validate_upload(raw_data, raw_filename, mime_type)
    except ValueError as e:
        upload_record = Upload(
            session_id=sid,
            filename=raw_filename,
            file_size=len(raw_data),
            file_type=Path(raw_filename).suffix.lower(),
            status="rejected",
            error=str(e),
        )
        try:
            db.session.add(upload_record)
            db.session.commit()
        except Exception:
            db.session.rollback()
        raise FileValidationError(str(e))

    file_bytes = io.BytesIO(raw_data)
    file_bytes.seek(0)
    file_bytes.name = raw_filename

    text, error = extract_text_from_file(file_bytes)
    if error:
        upload_record = Upload(
            session_id=sid,
            filename=raw_filename,
            file_size=len(raw_data),
            file_type=Path(raw_filename).suffix.lower(),
            status="extraction_failed",
            error=error,
        )
        try:
            db.session.add(upload_record)
            db.session.commit()
        except Exception:
            db.session.rollback()
        return jsonify({"error": error}), 422

    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[:MAX_DOCUMENT_CHARS]

    upload_record = Upload(
        session_id=sid,
        filename=raw_filename,
        file_size=len(raw_data),
        file_type=Path(raw_filename).suffix.lower(),
        status="processed",
    )
    try:
        db.session.add(upload_record)
        db.session.commit()
    except Exception:
        db.session.rollback()

    return jsonify({
        "text": text,
        "filename": raw_filename,
        "file_size": len(raw_data),
    })


# ---------------------------------------------------------------------------
# EXPORT
# ---------------------------------------------------------------------------


@app.route("/api/download-docx", methods=["POST"])
@export_rate_limit
def download_docx():
    data = request.get_json(silent=True)
    text = data.get("text", "") if data else ""
    if not text or not text.strip():
        raise ValidationError("No content provided.")

    text = str(text)
    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[:MAX_DOCUMENT_CHARS]

    filename = safe_filename(data.get("filename", ""), "document") + ".docx"

    sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"
    export_lock = _get_export_lock(f"docx:{sid}")

    with export_lock:
        export_record = Export(
            session_id=sid,
            fmt="docx",
            doc_type="document",
            status="started",
        )
        try:
            start = time.time()
            docx_bytes = create_word_doc(text)
            content_size = len(docx_bytes)
            duration_ms = int((time.time() - start) * 1000)
            export_record.status = "completed"
            export_record.success = True
            export_record.content_size = content_size
            try:
                db.session.add(export_record)
                db.session.commit()
            except Exception:
                db.session.rollback()
            return jsonify({
                "filename": filename,
                "content": docx_bytes.hex(),
                "size": content_size,
            })
        except Exception as e:
            export_record.status = "failed"
            export_record.error = str(e)
            try:
                db.session.add(export_record)
                db.session.commit()
            except Exception:
                db.session.rollback()
            raise ExportError(f"DOCX generation failed: {str(e)[:200]}")


@app.route("/api/download-pdf", methods=["POST"])
@export_rate_limit
def download_pdf():
    data = request.get_json(silent=True)
    text = data.get("text", "") if data else ""
    if not text or not text.strip():
        raise ValidationError("No content provided.")

    text = str(text)
    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[:MAX_DOCUMENT_CHARS]

    filename = safe_filename(data.get("filename", ""), "document") + ".pdf"

    sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"
    export_lock = _get_export_lock(f"pdf:{sid}")

    with export_lock:
        export_record = Export(
            session_id=sid,
            fmt="pdf",
            doc_type="document",
            status="started",
        )
        try:
            start = time.time()
            pdf_bytes = create_pdf(text)
            content_size = len(pdf_bytes)
            export_record.status = "completed"
            export_record.success = True
            export_record.content_size = content_size
            try:
                db.session.add(export_record)
                db.session.commit()
            except Exception:
                db.session.rollback()
            return jsonify({
                "filename": filename,
                "content": pdf_bytes.hex(),
                "size": content_size,
            })
        except Exception as e:
            export_record.status = "failed"
            export_record.error = str(e)
            try:
                db.session.add(export_record)
                db.session.commit()
            except Exception:
                db.session.rollback()
            raise ExportError(f"PDF generation failed: {str(e)[:200]}")


# ---------------------------------------------------------------------------
# GENERATE-RESUME (non-streaming)
# ---------------------------------------------------------------------------


DOC_SYSTEM_PROMPTS = {
    "resume": """You are a professional resume writer. Output ONLY valid JSON with this exact schema:
{
  "name": "Full Name",
  "contact": { "email": "...", "phone": "...", "location": "...", "linkedin": "...", "github": "..." },
  "summary": "2-3 sentence professional summary with **bold metrics**.",
  "experience": [{ "company": "", "role": "", "location": "", "dates": "", "bullets": ["Achievement **X%** ..."] }],
  "projects": [{ "name": "", "technologies": "", "dates": "", "bullets": ["..."] }],
  "education": [{ "institution": "", "degree": "", "location": "", "dates": "", "details": "" }],
  "skills": { "languages": "", "frameworks": "", "tools": "" }
}
RULES: **bold** for metrics. Achievement-oriented bullets. FAANG style. 1 page. ATS-friendly.""",

    "cover_letter": """You are a professional cover letter writer. Output ONLY valid JSON with this schema:
{
  "sender": { "name": "", "email": "", "phone": "", "address": "" },
  "recipient": { "name": "", "company": "", "address": "" },
  "date": "",
  "subject": "",
  "greeting": "",
  "body": ["Paragraph 1...", "Paragraph 2...", "Paragraph 3..."],
  "closing": "Sincerely,",
  "sender_title": ""
}
RULES: Professional tone. 3-4 paragraphs. Tailored to the role. Use **bold** for emphasis where appropriate.""",

    "proposal": """You are a business proposal writer. Output ONLY valid JSON with this schema:
{
  "title": "",
  "prepared_by": "",
  "prepared_for": "",
  "date": "",
  "executive_summary": "",
  "problem_statement": "",
  "solution": "",
  "scope": [{ "title": "", "description": "" }],
  "pricing": [{ "item": "", "cost": 0 }],
  "budget": "",
  "timeline": "",
  "deliverables": ["Deliverable 1", "Deliverable 2"]
}
RULES: Professional. Persuasive. Clear pricing. Use **bold** for key numbers.""",

    "report": """You are a professional report writer. Output ONLY valid JSON with this schema:
{
  "title": "",
  "prepared_by": "",
  "date": "",
  "abstract": "",
  "executive_summary": "",
  "introduction": "",
  "objectives": "",
  "methodology": "",
  "implementation": "",
  "analysis": "",
  "findings": "",
  "recommendations": "",
  "conclusion": "",
  "references": ["Reference 1", "Reference 2"]
}
RULES: Academic/professional tone. Use **bold** for key findings. Structured format.""",

    "invoice": """You are an invoice generator. Output ONLY valid JSON with this schema:
{
  "business": { "name": "", "email": "", "phone": "", "address": "" },
  "client": { "name": "", "company": "", "email": "", "address": "" },
  "invoice_number": "INV-001",
  "date": "",
  "due_date": "",
  "currency": "$",
  "items": [{ "description": "", "qty": 1, "rate": 0 }],
  "tax_rate": 0.18,
  "payment_method": "Bank Transfer",
  "notes": ""
}
RULES: Professional. Clean. Itemized. Calculate totals from qty * rate.""",

    "email": """You are a professional email writer. Output ONLY valid JSON with this schema:
{
  "to": "",
  "cc": "",
  "subject": "",
  "greeting": "Hi,",
  "body": ["Paragraph 1", "Paragraph 2"],
  "body_text": "",
  "closing": "Best regards,",
  "sender_name": "",
  "sender_title": "",
  "sender_email": ""
}
RULES: Use **bold** for emphasis. Tone varies by context (professional/friendly). Keep concise.""",

    "documentation": """You are a professional technical writer. Output ONLY valid JSON with this schema:
{
  "title": "",
  "author": "",
  "date": "",
  "version": "1.0",
  "overview": "Brief overview of the document",
  "sections": [
    {
      "heading": "Section Title",
      "body": "Section content with **bold** for key terms.",
      "subsections": [
        { "heading": "Subsection Title", "body": "Subsection content." }
      ],
      "code": "optional code block"
    }
  ],
  "conclusion": ""
}
RULES: Technical, clear, structured. Use **bold** for key terms and important concepts. Include code examples where relevant.""",

    "generic": """You are a **universal AI document writer** — you can generate ANY type of document the user requests. Output ONLY valid JSON with this flexible schema:
{
  "title": "Document Title",
  "author": "",
  "date": "",
  "document_type": "Specific document type (e.g. Meeting Minutes, NDA, Marketing Plan)",
  "summary": "Brief overview of the document",
  "sections": [
    {
      "heading": "Section Heading",
      "body": "Section content with **bold** for key terms and emphasis.",
      "items": ["Bullet point 1", "Bullet point 2"],
      "subsections": [
        { "heading": "Subsection", "body": "Details." }
      ]
    }
  ],
  "content": ["Paragraph 1", "Paragraph 2"],
  "conclusion": ""
}
RULES:
- Adapt to the SPECIFIC document type requested. Use correct professional structure for each type.
- For legal docs (NDA, contracts): include proper sections like parties, effective date, terms, signatures.
- For business docs (plans, strategies): include executive summary, market analysis, action items.
- For technical docs: include overview, architecture, implementation details, configuration.
- For meeting docs: include date, attendees, agenda items, discussion notes, action items.
- For academic docs: include abstract, methodology, findings, references.
- Use **bold** for emphasis on key terms, numbers, and headings.
- Professional tone throughout. Well-structured with clear sections."""
}


@app.route("/api/generate-resume", methods=["POST"])
@generation_rate_limit
def generate_resume():
    data = request.get_json(silent=True)
    if not data:
        raise ValidationError("No data provided.")

    try:
        prompt = validate_prompt(data.get("prompt", ""))
    except ValueError as e:
        raise ValidationError(str(e))

    doc_type = sanitize_text(data.get("doc_type", "generic"), max_length=50)
    system_prompt = DOC_SYSTEM_PROMPTS.get(doc_type, DOC_SYSTEM_PROMPTS["generic"])

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise AIError("API key not configured.")

    base_url = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
    client = OpenAI(api_key=api_key, base_url=base_url)

    start = time.time()
    sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"
    gen_key = build_generation_key(sid)

    if not burst_limiter.try_acquire(gen_key):
        return jsonify({"error": "A generation is already in progress."}), 429

    try:
        response = client.chat.completions.create(
            model=data.get("model", "deepseek-chat"),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt},
            ],
            max_tokens=4000,
        )
        content = response.choices[0].message.content or ""

        json_match = re.search(r"\{.*\}", content, re.DOTALL)
        if not json_match:
            raise AIError("AI response was not valid JSON.", 422)

        doc_data = json.loads(json_match.group())
        duration_ms = int((time.time() - start) * 1000)

        gen_record = Generation(
            session_id=sid,
            doc_type=doc_type,
            prompt_length=len(prompt),
            status="completed",
            success=True,
            model=data.get("model", "deepseek-chat"),
            duration_ms=duration_ms,
            tool_calls=0,
        )
        try:
            db.session.add(gen_record)
            db.session.commit()
        except Exception:
            db.session.rollback()

        return jsonify({"resume": doc_data})

    except json.JSONDecodeError:
        raise AIError("AI response contained malformed JSON.")
    except AIError:
        raise
    except Exception as e:
        duration_ms = int((time.time() - start) * 1000)
        gen_record = Generation(
            session_id=sid,
            doc_type=doc_type,
            prompt_length=len(prompt),
            status="failed",
            success=False,
            error=str(e)[:300],
            duration_ms=duration_ms,
        )
        try:
            db.session.add(gen_record)
            db.session.commit()
        except Exception:
            db.session.rollback()
        raise AIError(str(e)[:300])

    finally:
        burst_limiter.release(gen_key)


# ---------------------------------------------------------------------------
# STREAMING GENERATION
# ---------------------------------------------------------------------------


@app.route("/api/generate-stream", methods=["POST"])
@generation_rate_limit
def generate_stream():
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "No data provided"}), 400

    try:
        prompt = validate_prompt(data.get("prompt", ""))
    except ValueError as e:
        return jsonify({"error": str(e)}), 422

    doc_type = sanitize_text(data.get("doc_type", "generic"), max_length=50)
    system_prompt = DOC_SYSTEM_PROMPTS.get(doc_type, DOC_SYSTEM_PROMPTS["generic"])

    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        return jsonify({"error": "API key not configured"}), 500

    base_url = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com")
    client = OpenAI(api_key=api_key, base_url=base_url)

    sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"
    gen_key = build_generation_key(sid)

    if not burst_limiter.try_acquire(gen_key):
        return jsonify({"error": "A generation is already in progress."}), 429

    def event_stream():
        full_text = ""
        start_time = time.time()
        stream_active = True
        try:
            response = client.chat.completions.create(
                model=data.get("model", "deepseek-chat"),
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=4000,
                stream=True,
            )
            for chunk in response:
                if not stream_active:
                    break
                if time.time() - start_time > 120:
                    yield f"data: {json.dumps({'e': 'Generation timed out.'})}\n\n"
                    stream_active = False
                    break
                if chunk.choices and chunk.choices[0].delta and chunk.choices[0].delta.content:
                    token = chunk.choices[0].delta.content
                    full_text += token
                    yield f"data: {json.dumps({'t': token})}\n\n"

            if stream_active:
                json_match = re.search(r"\{.*\}", full_text, re.DOTALL)
                if json_match:
                    yield f"data: {json.dumps({'d': json.loads(json_match.group())})}\n\n"
                else:
                    yield f"data: {json.dumps({'e': 'Could not parse JSON from AI response', 'r': full_text[:500]})}\n\n"

        except Exception as e:
            error_msg = str(e)[:300]
            yield f"data: {json.dumps({'e': error_msg})}\n\n"
        finally:
            burst_limiter.release(gen_key)
            gen_record = Generation(
                session_id=sid,
                doc_type=doc_type,
                prompt_length=len(prompt),
                status="completed" if stream_active and full_text else "failed",
                success=bool(full_text),
                error=None if stream_active else "interrupted",
                duration_ms=int((time.time() - start_time) * 1000),
            )
            try:
                with app.app_context():
                    db.session.add(gen_record)
                    db.session.commit()
            except Exception:
                pass

    resp = Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
    )
    resp.headers["Cache-Control"] = "no-cache"
    resp.headers["X-Accel-Buffering"] = "no"
    resp.headers["X-Content-Type-Options"] = "nosniff"
    return resp


# ---------------------------------------------------------------------------
# GENERATE HTML TO DOCX
# ---------------------------------------------------------------------------


@app.route("/api/generate-html-docx", methods=["POST"])
@export_rate_limit
def generate_html_docx():
    from bs4 import BeautifulSoup

    data = request.get_json(silent=True)
    if not data:
        raise ValidationError("No data provided.")

    html_content = sanitize_text((data or {}).get("html", ""), max_length=500000)
    css_content = (data or {}).get("css", "")
    filename = safe_filename((data or {}).get("filename", ""), "resume") + ".docx"

    if not html_content.strip():
        raise ValidationError("No HTML provided.")

    try:
        doc = docx.Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        pf = style.paragraph_format
        pf.space_after = Pt(3)
        pf.space_before = Pt(0)
        pf.line_spacing = 1.15

        def add_run(paragraph, text, bold=False, italic=False, underline=False, size=None, color=None, font_name=None):
            if not text or not text.strip():
                return
            safe_text = sanitize_text(text)
            run = paragraph.add_run(safe_text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if underline:
                run.underline = True
            if size:
                run.font.size = size
            if color:
                run.font.color.rgb = color
            if font_name:
                run.font.name = font_name
            return run

        def process_inline(paragraph, parent):
            for child in parent.children:
                if isinstance(child, str):
                    add_run(paragraph, child)
                elif child.name in ('strong', 'b'):
                    add_run(paragraph, child.get_text(), bold=True)
                elif child.name in ('em', 'i'):
                    add_run(paragraph, child.get_text(), italic=True)
                elif child.name == 'br':
                    add_run(paragraph, '\n')
                elif child.name == 'span':
                    style_attr = child.get('style', '')
                    bold = 'font-weight' in style_attr and ('bold' in style_attr or '700' in style_attr)
                    add_run(paragraph, child.get_text(), bold=bold)
                elif child.name == 'a':
                    add_run(paragraph, child.get_text(), underline=True)
                else:
                    process_inline(paragraph, child)

        def has_text_content(el):
            text = el.get_text(strip=True)
            return bool(text)

        soup = BeautifulSoup(html_content, 'html.parser')

        body = soup.find('body')
        root = body if body else soup

        for el in root.find_all(recursive=False):
            if not has_text_content(el):
                continue
            tag = el.name if hasattr(el, 'name') else None
            if not tag or tag in ('script', 'style', 'nav', 'aside', 'meta', 'link'):
                continue

            if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if tag == 'h1' else WD_ALIGN_PARAGRAPH.LEFT
                level = int(tag[1])
                sizes = {1: Pt(20), 2: Pt(14), 3: Pt(12), 4: Pt(11)}
                size = sizes.get(level, Pt(11))
                text = el.get_text(strip=True)
                add_run(p, text, bold=True, size=size)
                p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
                p.paragraph_format.space_after = Pt(4 if level <= 2 else 3)

            elif tag == 'p':
                p = doc.add_paragraph()
                process_inline(p, el)
                if p.text.strip():
                    p.paragraph_format.space_after = Pt(4)

            elif tag == 'div':
                cls = ' '.join(el.get('class', []))
                text = el.get_text(strip=True)
                if not text:
                    continue

                is_section = any(c in cls for c in ['r-section', 'prop-section', 'rpt-section', 'inv-header',
                                                     'section-label', 'rpt-abstract-label', 'cl-subject',
                                                     'email-subject-line'])
                is_header = any(c in cls for c in ['r-exp-header', 'r-edu-header', 'email-header-bar',
                                                    'prop-card-title', 'cl-sender', 'cl-name', 'cl-greeting',
                                                    'cl-closing', 'inv-from', 'inv-company', 'inv-number',
                                                    'prop-cover-page'])
                is_body = any(c in cls for c in ['cl-body', 'email-body-content', 'prop-section', 'rpt-section'])

                if is_section:
                    p = doc.add_paragraph()
                    add_run(p, text, bold=True, size=Pt(10))
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(3)
                elif is_header:
                    p = doc.add_paragraph()
                    for child in el.children:
                        if isinstance(child, str):
                            add_run(p, child)
                        elif child.name in ('strong', 'b', 'span'):
                            is_b = child.name in ('strong', 'b') or ('font-weight' in (child.get('style','')) and 'bold' in child.get('style',''))
                            add_run(p, child.get_text(), bold=is_b)
                        else:
                            add_run(p, child.get_text())
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(2)
                elif is_body:
                    for sub_el in el.find_all(['p', 'div'], recursive=False):
                        sub_text = sub_el.get_text(strip=True)
                        if sub_text:
                            p = doc.add_paragraph()
                            process_inline(p, sub_el)
                            p.paragraph_format.space_after = Pt(4)
                else:
                    p = doc.add_paragraph()
                    process_inline(p, el)
                    p.paragraph_format.space_after = Pt(3)

            elif tag == 'ul':
                for li in el.find_all('li', recursive=False):
                    text = li.get_text(strip=True)
                    if text:
                        p = doc.add_paragraph(style='List Bullet')
                        process_inline(p, li)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.3)

            elif tag == 'ol':
                for i, li in enumerate(el.find_all('li', recursive=False)):
                    text = li.get_text(strip=True)
                    if text:
                        p = doc.add_paragraph()
                        add_run(p, f"{i+1}. ", bold=True)
                        process_inline(p, li)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.3)

            elif tag == 'hr':
                p = doc.add_paragraph()
                add_run(p, '─' * 65, size=Pt(6))
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

            elif tag == 'blockquote':
                p = doc.add_paragraph()
                process_inline(p, el)
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

            elif tag == 'table':
                for tr in el.find_all('tr'):
                    cells = tr.find_all(['td', 'th'])
                    if cells:
                        row_text = ' | '.join(c.get_text(strip=True) for c in cells)
                        p = doc.add_paragraph()
                        add_run(p, row_text, size=Pt(9.5))
                        p.paragraph_format.space_after = Pt(2)

            else:
                text = el.get_text(strip=True)
                if text:
                    p = doc.add_paragraph()
                    process_inline(p, el)
                    p.paragraph_format.space_after = Pt(3)

        buffer = io.BytesIO()
        doc.save(buffer)

        sid = request.headers.get("X-Session-Id", "") or request.remote_addr or "unknown"
        export_record = Export(
            session_id=sid,
            fmt="html-docx",
            status="completed",
            success=True,
            content_size=buffer.tell(),
        )
        try:
            db.session.add(export_record)
            db.session.commit()
        except Exception:
            db.session.rollback()

        return jsonify({"filename": filename, "content": buffer.getvalue().hex()})
    except Exception as e:
        raise ExportError(f"HTML to DOCX conversion failed: {str(e)[:200]}")


# ---------------------------------------------------------------------------
# NEW SESSION
# ---------------------------------------------------------------------------


@app.route("/api/new-session", methods=["POST"])
def new_session():
    sid = str(uuid.uuid4())
    with _sessions_lock:
        sessions[sid] = {
            "memory": Memory(),
            "thinking": [],
            "last_doc_hash": None,
            "doc_loaded": False,
            "last_download_name": "document",
            "created_at": time.time(),
        }
    return jsonify({"session_id": sid})


# ---------------------------------------------------------------------------
# HEALTH CHECK
# ---------------------------------------------------------------------------


@app.route("/api/health", methods=["GET"])
def health_check():
    return jsonify({
        "status": "ok",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "sessions_active": len(sessions),
    })


# ---------------------------------------------------------------------------
# SOCKETIO
# ---------------------------------------------------------------------------


@socketio.on("connect")
def handle_connect():
    emit("session_id", {"session_id": request.sid})


@socketio.on("disconnect")
def handle_disconnect():
    gen_key = build_generation_key(request.sid)
    burst_limiter.release(gen_key)


@socketio.on("send_message")
def handle_socket_message(data):
    sid = request.sid
    raw_input = data.get("message", "").strip() if data else ""
    raw_doc_text = (data or {}).get("document_text", "").strip()
    doc_filename = (data or {}).get("document_filename", DEFAULT_FILENAME)

    if not raw_input:
        emit("agent_response", {
            "answer": "",
            "thinking": [],
            "success": False,
        })
        return

    try:
        user_input = sanitize_text(strip_html_tags(raw_input), max_length=5000)
    except Exception:
        emit("agent_response", {
            "answer": "Invalid message.",
            "thinking": [],
            "success": False,
        })
        return

    doc_text = ""
    if raw_doc_text:
        try:
            doc_text = validate_document_text(raw_doc_text)
        except ValueError:
            pass

    doc_filename = sanitize_filename(doc_filename) or DEFAULT_FILENAME

    sess = get_session(sid)
    current_hash = document_hash(doc_text) if doc_text else None
    if current_hash and sess["last_doc_hash"] != current_hash:
        sess["doc_loaded"] = False

    should_mark_doc = bool(doc_text and not sess["doc_loaded"])
    task = f"[DOCUMENT: {doc_filename}]\n{doc_text}\n\nRequest: {user_input}" if should_mark_doc else user_input

    sess["thinking"] = []

    def on_tool_call(tool_name, tool_input, tool_output):
        try:
            emit("thinking_step", {
                "tool": tool_name,
                "input": tool_input,
                "output": str(tool_output)[:300],
            })
        except Exception:
            pass

    gen_key = build_generation_key(sid)
    if not burst_limiter.try_acquire(gen_key):
        emit("agent_response", {
            "answer": "Please wait for the current generation to complete.",
            "thinking": [],
            "success": False,
        })
        return

    try:
        result = safe_agent_call(
            task=task,
            config=DOCUMENT_AGENT_CONFIG,
            memory=sess["memory"],
            on_tool_call=on_tool_call,
        )

        answer = result.get("answer", "Sorry, I could not complete that request.")
        success = result.get("success", False)

        if should_mark_doc and success:
            sess["last_doc_hash"] = current_hash
            sess["doc_loaded"] = True
        sess["last_download_name"] = doc_filename if doc_text else "document_latest"

        emit("agent_response", {
            "answer": answer,
            "thinking": sess["thinking"],
            "success": success,
        })
    finally:
        burst_limiter.release(gen_key)


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------


if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    debug = os.getenv("FLASK_DEBUG", "0") == "1"
    host = os.getenv("HOST", "0.0.0.0")
    logger.info("Starting Document Agent on %s:%s (debug=%s)", host, port, debug)
    socketio.run(app, host=host, port=port, debug=debug)
